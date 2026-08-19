from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "flow_assets"
OUTPUT = DOCS / "Imdad_AI_Project_Flowchart_Document.docx"

BLUE = "#0C66EE"
BLUE_2 = "#468EF9"
CYAN = "#18C9FF"
SOFT_BLUE = "#EAF3FF"
INK = "#1F2937"
MUTED = "#666666"
BORDER = "#BFD7FF"
GREEN = "#18B978"
AMBER = "#F2B94B"
RED = "#E55067"
WHITE = "#FFFFFF"


def font(size=28, bold=False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def rounded_box(draw, xy, text, fill=WHITE, outline=BORDER, text_color=INK, radius=22, width=3, title=False):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)
    max_chars = max(14, int((x2 - x1) / 17))
    lines = []
    for part in text.split("\n"):
        lines.extend(wrap(part, max_chars) or [""])
    fnt = font(30 if title else 25, bold=title)
    line_h = int(fnt.size * 1.25)
    total_h = line_h * len(lines)
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        draw.text((x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2, y), line, fill=text_color, font=fnt)
        y += line_h


def diamond(draw, cx, cy, w, h, text, fill=WHITE, outline=BORDER):
    points = [(cx, cy - h // 2), (cx + w // 2, cy), (cx, cy + h // 2), (cx - w // 2, cy)]
    draw.polygon(points, fill=fill, outline=outline)
    draw.line(points + [points[0]], fill=outline, width=4)
    lines = wrap(text, 18)
    fnt = font(24, bold=True)
    line_h = 30
    y = cy - (line_h * len(lines)) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        draw.text((cx - (bbox[2] - bbox[0]) / 2, y), line, fill=INK, font=fnt)
        y += line_h


def arrow(draw, start, end, color=BLUE, width=5):
    draw.line([start, end], fill=color, width=width)
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) > abs(ey - sy):
      direction = 1 if ex > sx else -1
      head = [(ex, ey), (ex - 18 * direction, ey - 10), (ex - 18 * direction, ey + 10)]
    else:
      direction = 1 if ey > sy else -1
      head = [(ex, ey), (ex - 10, ey - 18 * direction), (ex + 10, ey - 18 * direction)]
    draw.polygon(head, fill=color)


def canvas(path, size=(1600, 1080), title=""):
    img = Image.new("RGB", size, "#F7FBFF")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((30, 30, size[0] - 30, size[1] - 30), radius=34, fill=WHITE, outline="#DDEBFF", width=3)
    if title:
        draw.text((70, 58), title, fill=INK, font=font(42, bold=True))
        draw.line((70, 120, size[0] - 70, 120), fill=SOFT_BLUE, width=5)
    return img, draw


def system_flow():
    path = ASSETS / "system_flow.png"
    img, draw = canvas(path, title="Imdad AI - End-to-End System Flow")

    rounded_box(draw, (80, 170, 360, 270), "User opens\nImdad AI", fill=SOFT_BLUE, title=True)
    rounded_box(draw, (450, 170, 730, 270), "Signup / Login\nwith role", fill=SOFT_BLUE, title=True)
    diamond(draw, 900, 220, 230, 150, "Role?")
    rounded_box(draw, (1110, 130, 1450, 215), "User\nAI Help workspace", fill="#ECFDF5", outline="#B7E9CE", title=True)
    rounded_box(draw, (1110, 245, 1450, 330), "Support Engineer\nTicket dashboard", fill="#FFF8E7", outline="#F4D48A", title=True)
    rounded_box(draw, (1110, 360, 1450, 445), "Admin\nAnalytics control", fill="#F4ECFF", outline="#D5B9FF", title=True)
    arrow(draw, (360, 220), (450, 220))
    arrow(draw, (730, 220), (785, 220))
    arrow(draw, (1015, 205), (1110, 172))
    arrow(draw, (1015, 220), (1110, 288))
    arrow(draw, (1010, 245), (1110, 402))

    rounded_box(draw, (80, 575, 325, 670), "User enters\ntechnical issue", title=True)
    rounded_box(draw, (395, 575, 640, 670), "Search\nKnowledge Base", fill=SOFT_BLUE, title=True)
    diamond(draw, 785, 622, 220, 150, "KB match?")
    rounded_box(draw, (955, 530, 1225, 625), "Show KB solution", fill="#ECFDF5", outline="#B7E9CE", title=True)
    rounded_box(draw, (955, 690, 1225, 785), "Generate AI\nsuggestion", fill="#EAF8FF", outline="#A9E9FF", title=True)
    diamond(draw, 1375, 622, 220, 150, "Resolved?")
    rounded_box(draw, (1215, 850, 1510, 940), "Auto-create ticket\nif unresolved", fill="#FFF8E7", outline="#F4D48A", title=True)
    rounded_box(draw, (80, 850, 360, 940), "Save chat history\nand result", fill="#ECFDF5", outline="#B7E9CE", title=True)
    arrow(draw, (325, 622), (395, 622))
    arrow(draw, (640, 622), (675, 622))
    arrow(draw, (895, 585), (955, 578))
    arrow(draw, (895, 662), (955, 735))
    arrow(draw, (1225, 578), (1265, 600))
    arrow(draw, (1225, 735), (1265, 655))
    arrow(draw, (1375, 697), (1375, 850))
    arrow(draw, (1265, 622), (360, 892), color=GREEN, width=4)
    img.save(path)
    return path


def ticket_pipeline():
    path = ASSETS / "ticket_pipeline.png"
    img, draw = canvas(path, size=(1600, 900), title="Ticket Creation Pipeline")
    boxes = [
        ((70, 190, 300, 300), "Unresolved\nissue", RED),
        ((360, 190, 590, 300), "Generate\nTicket ID", BLUE),
        ((650, 190, 880, 300), "Save user\nand issue", BLUE),
        ((940, 190, 1170, 300), "Save AI\nsuggestion", BLUE),
        ((1230, 190, 1530, 300), "Set status\nOpen", GREEN),
        ((200, 520, 480, 640), "Detect priority\nLow / Medium /\nHigh / Critical", AMBER),
        ((560, 520, 840, 640), "Assign team\nby issue type", CYAN),
        ((920, 520, 1200, 640), "Store timestamps\nCreated / Updated", BLUE_2),
        ((1280, 520, 1530, 640), "Visible to User,\nEngineer, Admin", GREEN),
    ]
    for xy, label, color in boxes:
        rounded_box(draw, xy, label, fill=WHITE, outline=color, title=True)
    for a, b in [((300, 245), (360, 245)), ((590, 245), (650, 245)), ((880, 245), (940, 245)), ((1170, 245), (1230, 245))]:
        arrow(draw, a, b)
    arrow(draw, (1380, 300), (340, 520))
    for a, b in [((480, 580), (560, 580)), ((840, 580), (920, 580)), ((1200, 580), (1280, 580))]:
        arrow(draw, a, b)
    img.save(path)
    return path


def role_flow():
    path = ASSETS / "role_flow.png"
    img, draw = canvas(path, size=(1600, 980), title="Role-Based Workflow")
    lanes = [
        (150, "User", "#ECFDF5", "#B7E9CE", ["Describe issue", "Get KB or AI suggestion", "Resolve or raise ticket", "Track own tickets"]),
        (420, "Support Engineer", "#EAF8FF", "#A9E9FF", ["View assigned tickets", "Inspect AI context", "Assign and update status", "Resolve or close ticket"]),
        (690, "Admin", "#F4ECFF", "#D5B9FF", ["View analytics", "Manage users and engineers", "Manage tickets", "Manage Knowledge Base"]),
    ]
    for y, title, fill, outline, steps in lanes:
        rounded_box(draw, (70, y, 305, y + 90), title, fill=fill, outline=outline, title=True)
        x = 400
        previous = (305, y + 45)
        for step in steps:
            rounded_box(draw, (x, y, x + 245, y + 90), step, fill=WHITE, outline=outline, title=True)
            arrow(draw, previous, (x, y + 45))
            previous = (x + 245, y + 45)
            x += 285
    img.save(path)
    return path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(31, 41, 55)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color in [("Heading 1", 16, "0C66EE"), ("Heading 2", 13, "0C66EE"), ("Heading 3", 12, "1F4D78")]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = "Imdad AI | Project Flow and Pipeline"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(102, 102, 102)
    footer = section.footer.paragraphs[0]
    footer.text = "Client presentation document"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(102, 102, 102)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Imdad AI")
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = RGBColor(31, 41, 55)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(16)
    r2 = p2.add_run("Project Flow and Pipeline Overview")
    r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor(12, 102, 238)

    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    for idx, (label, value) in enumerate([
        ("Product", "AI-powered IT ticket desk"),
        ("Roles", "User, Support Engineer, Admin"),
        ("Purpose", "Resolve faster and track support work"),
    ]):
        cell = table.cell(0, idx)
        set_cell_shading(cell, "EAF3FF")
        set_cell_text(cell, f"{label}\n{value}", bold=True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def build_doc():
    ASSETS.mkdir(parents=True, exist_ok=True)
    system = system_flow()
    ticket = ticket_pipeline()
    roles = role_flow()

    doc = Document()
    style_doc(doc)
    add_title(doc)

    doc.add_heading("1. Project Summary", level=1)
    doc.add_paragraph(
        "Imdad AI is a role-based AI help desk system that helps users report technical issues, receive instant Knowledge Base or AI troubleshooting, and create support tickets when human support is required."
    )
    add_bullets(
        doc,
        [
            "Frontend: React and Vite with a premium SaaS-style multi-page UI.",
            "Backend/data layer: Supabase for tickets, Knowledge Base entries, and chat history.",
            "AI workflow: Knowledge Base-first search, then AI troubleshooting when no KB match is found.",
            "Roles: User, Support Engineer, and Admin, each with its own dashboard and permissions.",
        ],
    )

    doc.add_heading("2. End-to-End System Flow", level=1)
    doc.add_paragraph("The full system starts with role-based access, then routes each role into the right workspace.")
    doc.add_picture(str(system), width=Inches(6.7))

    doc.add_heading("3. User Support Pipeline", level=1)
    add_bullets(
        doc,
        [
            "The user enters a technical issue in the AI Help module.",
            "The system searches Knowledge Base entries by keyword.",
            "If a match exists, the KB solution is shown immediately.",
            "If no match exists, the issue is sent to the AI troubleshooting layer.",
            "If the issue is resolved, the interaction is stored as resolved chat history.",
            "If unresolved, a ticket is created automatically and becomes trackable.",
        ],
    )

    doc.add_heading("4. Ticket Creation Pipeline", level=1)
    doc.add_paragraph("When human support is needed, the system converts the conversation into a structured support ticket.")
    doc.add_picture(str(ticket), width=Inches(6.7))

    doc.add_heading("5. Role-Based Workflow", level=1)
    doc.add_paragraph("Each role has a focused workflow so users only see what they need and support staff can manage work efficiently.")
    doc.add_picture(str(roles), width=Inches(6.7))

    doc.add_heading("6. Client Value", level=1)
    add_bullets(
        doc,
        [
            "Reduces repetitive support work by solving common IT issues through the Knowledge Base.",
            "Improves response speed by using AI suggestions before escalation.",
            "Creates consistent, complete tickets for unresolved issues.",
            "Gives engineers the context needed to resolve tickets faster.",
            "Gives admins visibility into ticket status, priority, workload, and support performance.",
        ],
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_doc()
