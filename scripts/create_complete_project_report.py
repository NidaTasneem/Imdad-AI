from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Imdad_AI_Complete_Project_Report.docx"
LOGO = ROOT / "public" / "brand" / "imdad-ai-mark.png"
ASSET_DIR = ROOT / "docs" / "report_assets"
AUTHOR_NAME = "Nida Tasneem"

BLACK = RGBColor(20, 24, 32)
DARK = RGBColor(35, 42, 56)
MUTED = RGBColor(92, 99, 112)
LINE = "D7DCE5"
FILL = "F3F5F8"
FILL_DARK = "E7EBF1"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = LINE, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_row_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def set_cell_margins(cell, top=105, start=210, bottom=105, end=210) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def paragraph(text: str, style: str | None = None, bold: bool = False, size: float | None = None, color=BLACK):
    p = DOC.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return p


def add_kicker(text: str) -> None:
    p = DOC.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = MUTED
    r.font.name = "Calibri"


def add_h1(text: str) -> None:
    p = DOC.add_paragraph(text, style="Heading 1")
    p.paragraph_format.keep_with_next = True


def add_h2(text: str) -> None:
    p = DOC.add_paragraph(text, style="Heading 2")
    p.paragraph_format.keep_with_next = True


def add_body(text: str) -> None:
    p = DOC.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10


def add_bullets(items: list[str]) -> None:
    for item in items:
        p = DOC.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.color.rgb = BLACK


def add_numbered(items: list[str]) -> None:
    for item in items:
        p = DOC.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.color.rgb = BLACK


def add_callout(title: str, body: str) -> None:
    table = DOC.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(6.35)
    set_row_cant_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, FILL)
    set_cell_border(cell)
    set_cell_margins(cell, top=170, bottom=170, start=230, end=230)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = BLACK
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = DARK
    DOC.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = DOC.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_row_repeat_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        if widths:
            hdr_cells[i].width = Inches(widths[i])
        set_cell_shading(hdr_cells[i], FILL_DARK)
        set_cell_border(hdr_cells[i])
        set_cell_margins(hdr_cells[i], top=105, bottom=105, start=220, end=220)
        p = hdr_cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9.3)
        run.font.color.rgb = BLACK
    for row in rows:
        row_obj = table.add_row()
        set_row_cant_split(row_obj)
        cells = row_obj.cells
        for i, value in enumerate(row):
            if widths:
                cells[i].width = Inches(widths[i])
            set_cell_border(cells[i])
            set_cell_margins(cells[i], top=95, bottom=95, start=220, end=220)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            run.font.size = Pt(9.2)
            run.font.color.rgb = DARK
    DOC.add_paragraph().paragraph_format.space_after = Pt(2)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_separate)

    result = OxmlElement("w:t")
    result.text = "1"
    run._r.append(result)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def add_picture_block(image_path: Path, caption: str, width: float = 6.35) -> None:
    if not image_path.exists():
        return
    p = DOC.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(image_path), width=Inches(width))
    cap = DOC.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_two_col_table(rows: list[tuple[str, str]]) -> None:
    add_table(["Area", "Details"], [[a, b] for a, b in rows], widths=[1.65, 4.85])


def set_footer(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""
    r = p.add_run(f"Imdad AI Complete Project Report | {AUTHOR_NAME} | Page ")
    r.font.size = Pt(8.5)
    r.font.color.rgb = MUTED
    add_field(p, "PAGE")
    r = p.add_run(" of ")
    r.font.size = Pt(8.5)
    r.font.color.rgb = MUTED
    add_field(p, "NUMPAGES")


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.35)
    set_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.8)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, before, after in [
        ("Heading 1", 16, 14, 7),
        ("Heading 2", 12.5, 10, 5),
        ("Heading 3", 11.3, 7, 3),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    return doc


DOC = setup_document()


def cover_page() -> None:
    if LOGO.exists():
        p = DOC.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(1.0))

    p = DOC.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("IMDAD AI")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = BLACK

    p = DOC.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI-Powered Ticket Management Agent")
    r.font.size = Pt(15)
    r.font.color.rgb = DARK

    p = DOC.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Complete Project Report and Mentor Presentation Documentation")
    r.font.size = Pt(11.5)
    r.font.color.rgb = MUTED

    table = DOC.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        set_row_cant_split(row)
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell, top=145, bottom=145, start=210, end=210)
    data = [
        ("Project Type", "Role-based AI help desk and support ticket workflow system"),
        ("Prepared For", "Mentor review, viva, demo presentation, and project evaluation"),
        ("Prepared By", AUTHOR_NAME),
        ("Current Status", "Core user, engineer, admin, database, analytics, and demo reset milestones completed"),
    ]
    for row, (label, value) in zip(table.rows, data):
        set_cell_shading(row.cells[0], FILL_DARK)
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(value)

    DOC.add_paragraph()
    add_callout(
        "Executive positioning",
        "Imdad AI demonstrates how a modern support desk can reduce repetitive IT tickets by combining Knowledge Base search, AI troubleshooting, automatic ticket creation, role-based dashboards, and live Supabase-backed workflow tracking.",
    )
    DOC.add_page_break()


def contents_page() -> None:
    p = DOC.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Table of Contents")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = BLACK
    toc = DOC.add_paragraph()
    toc.paragraph_format.space_after = Pt(0)
    add_field(toc, 'TOC \\o "1-1" \\h \\z \\u')
    DOC.add_page_break()


def executive_summary() -> None:
    add_h1("1. Executive Summary")
    add_body(
        "Imdad AI is a role-based AI help desk application built to help users solve common technical problems quickly and create structured support tickets only when human support is needed. The system supports Users, Support Engineers, and Admins, each with a focused dashboard and permission-aware navigation."
    )
    add_body(
        "The project started as an AI Ticket Management Agent concept and has evolved into a complete multi-page web application connected to Supabase. It now includes authentication, session management, AI help flow, Knowledge Base search, ticket creation, user ticket tracking, engineer ticket management, admin analytics, Knowledge Base management, and resolution history."
    )
    add_callout(
        "Why this project stands out",
        "The application does not only collect tickets. It tries to prevent unnecessary tickets first through Knowledge Base matching and AI troubleshooting, then creates a clean, trackable ticket with priority, team assignment, AI suggestion, and support workflow metadata.",
    )
    add_two_col_table(
        [
            ("Project name", "Imdad AI"),
            ("Core idea", "Smart AI-powered support ticket system"),
            ("Primary users", "Employees, support engineers, and administrators"),
            ("Database status", "Supabase schema repaired, demo data reset, and resolution history added"),
            ("Current milestone", "Modules through Reports and Analytics have been built and verified at prototype level"),
        ]
    )


def problem_objectives() -> None:
    add_h1("2. Problem Statement and Project Objectives")
    add_h2("Problem Statement")
    add_body(
        "In many organizations, users raise tickets for repeated issues such as VPN problems, password reset, email blocking, access denied, slow systems, printer issues, or application errors. Support engineers lose time answering the same issues manually, while admins have limited visibility into ticket trends, workload, and resolution performance."
    )
    add_h2("Objectives")
    add_bullets(
        [
            "Allow users to sign up, log in, describe technical issues, and receive instant help.",
            "Search the Knowledge Base before escalating to AI or human support.",
            "Automatically create support tickets when the user confirms that the issue is not resolved.",
            "Generate ticket ID, priority, team assignment, status, and timestamps automatically.",
            "Provide support engineers with ticket management, resolution notes, status updates, and history.",
            "Provide admins with user control, Knowledge Base management, ticket monitoring, and analytics.",
            "Store live data in Supabase so the prototype behaves like a real support system.",
        ]
    )
    add_h2("Scope")
    add_body(
        "The current scope is a working full-stack prototype suitable for mentor demonstration. It includes a React front end, Express AI API, Supabase database, role-based workflow, and a responsive SaaS-style interface. Production deployment, advanced Supabase Auth policies, audit-grade security, and notification integrations are listed as future enhancements."
    )


def journey() -> None:
    add_h1("3. Completed Project Journey")
    add_table(
        ["Phase", "Work Completed", "Result"],
        [
            ["Concept and requirements", "Defined User, Support Engineer, and Admin roles with AI help, ticketing, Knowledge Base, chat history, and analytics requirements.", "Clear role-based product scope."],
            ["Initial web app", "Built the React/Vite project and organized content into actual pages instead of placing everything on the home page.", "Professional multi-page app structure."],
            ["Interface upgrade", "Created a SaaS-style home page, role-based navbar, 3D assets, FAQ, footer, smooth transitions, and responsive layouts.", "The app gained a clearer visual hierarchy and improved presentation quality."],
            ["Supabase connection", "Connected the app to Supabase using environment variables and the supabase-js client.", "Real database reads and writes became available."],
            ["Database repair", "Fixed schema mismatches, removed old auth.users dependency, added password_hash, repaired ticket fields, and reset demo data.", "Clean Supabase schema with stable demo records."],
            ["Engineer modules", "Implemented engineer authentication, protected dashboard, open ticket management, update status, assign engineer, remarks, resolution notes, and resolution history.", "Support workflow became operational."],
            ["Admin and analytics", "Polished admin dashboard with KPIs, statistics, category/priority/status reporting, engineer performance, and Knowledge Base management.", "Admin can monitor complete system activity."],
            ["Testing and verification", "Verified build, login, ticket creation, engineer updates, resolution history, admin analytics, and Supabase refresh behavior.", "Current milestone is functionally complete for demo use."],
        ],
        widths=[1.35, 3.15, 2.0],
    )


def tech_stack() -> None:
    add_h1("4. Technology Stack")
    add_table(
        ["Layer", "Technology", "Purpose"],
        [
            ["Frontend", "React 19, TypeScript, Vite", "Fast single-page application with route-based views and reusable components."],
            ["Styling", "CSS, custom design system, 3D SVG/PNG assets, lucide-react icons", "Consistent SaaS-style interface, cards, dashboards, forms, tables, and visual hierarchy."],
            ["Backend API", "Node.js, Express", "Provides `/api/troubleshoot` for AI troubleshooting requests."],
            ["AI", "Ollama Llama 3.2 with fallback suggestions", "Generates technical troubleshooting steps when no Knowledge Base match exists."],
            ["Database", "Supabase Postgres", "Stores users, tickets, Knowledge Base entries, chat history, and resolution history."],
            ["Auth prototype", "Role-based app login with hashed password values and local session persistence", "Controls redirection and dashboard access by role."],
            ["Verification", "TypeScript build and browser testing", "Checks UI, workflows, persistence, and realtime-like data refresh."],
        ],
        widths=[1.2, 2.0, 3.3],
    )
    add_callout(
        "Important security note",
        "The current project is a prototype. Supabase Row Level Security policies are permissive for demo speed. In production, the next step is to replace prototype policies with authenticated Supabase Auth rules and role-specific policies.",
    )


def architecture() -> None:
    add_h1("5. System Architecture")
    add_body(
        "The system uses a client-first web architecture. React manages the interface, role navigation, forms, dashboards, and local session state. Supabase stores live application data. The Express server connects the AI Help module to Ollama Llama 3.2 and returns structured troubleshooting advice with a safe fallback when the local AI model is unavailable."
    )
    add_table(
        ["Component", "Responsibility"],
        [
            ["React application", "Displays pages, handles role-based navigation, validates inputs, calls Supabase, and renders dashboards."],
            ["Supabase client", "Reads and writes users, tickets, Knowledge Base entries, chat history, and ticket resolution history."],
            ["Express API", "Receives issue descriptions and requests AI-generated troubleshooting from Ollama."],
            ["Ollama Llama 3.2", "Generates troubleshooting steps for custom issues that are not solved by the Knowledge Base."],
            ["Local session", "Stores the current signed-in role and redirects users to the correct dashboard."],
        ],
        widths=[1.7, 4.8],
    )
    add_picture_block(
        ASSET_DIR / "system_architecture_diagram.png",
        "Figure 1. System architecture showing React, Supabase, Express API, Ollama, ticket workflow, and role dashboards.",
    )
    add_h2("Architecture Flow")
    add_numbered(
        [
            "User opens Imdad AI and selects the correct role.",
            "After login, the app stores a session and redirects based on the role.",
            "User enters an issue in AI Help.",
            "System searches Supabase Knowledge Base by keywords.",
            "If a solution exists, it is shown immediately.",
            "If no match exists, the Express API asks Ollama Llama 3.2 for troubleshooting steps.",
            "If resolved, chat history is saved.",
            "If not resolved, a ticket is created and stored in Supabase.",
            "Engineers update the ticket status, assignment, remarks, resolution notes, and history.",
            "Admin monitors users, tickets, Knowledge Base, statistics, and reports.",
        ]
    )


def role_modules() -> None:
    add_h1("6. Role-Based Modules")
    add_h2("User Role")
    add_bullets(
        [
            "Create account and log in.",
            "Use AI Help for troubleshooting.",
            "Receive Knowledge Base result or AI suggestion.",
            "Mark issue as resolved or proceed with human support.",
            "Create support tickets when needed.",
            "View only own tickets, own ticket details, own chat history, and resolution updates.",
        ]
    )
    add_h2("Support Engineer Role")
    add_bullets(
        [
            "Log in to a protected engineer dashboard.",
            "View open, assigned, and available support tickets.",
            "Search and filter tickets by issue, status, priority, team, and recency.",
            "Review user issue details, AI suggestion, and ticket context.",
            "Assign tickets, update status, add remarks, add resolution notes, resolve, and close tickets.",
            "Track resolution history linked to each ticket.",
        ]
    )
    add_h2("Admin Role")
    add_bullets(
        [
            "Log in to the admin dashboard.",
            "View users, support engineers, tickets, status reports, priority reports, category analysis, and engineer performance.",
            "Manage Knowledge Base entries with add, update, delete, and category selection.",
            "Monitor support activity, ticket volume, open workload, resolved tickets, and high-priority issues.",
        ]
    )
    add_table(
        ["Role", "Home after Login", "Key Protected Pages"],
        [
            ["User", "AI Help / user home", "AI Help, Raise Ticket, My Tickets, Ticket Details, Chat History, Knowledge Base"],
            ["Support Engineer", "Engineer Analytics Dashboard", "Engineer Home, Tickets, Ticket Details, Chat History, Knowledge Base"],
            ["Admin", "Admin Analytics Dashboard", "Admin Home, Users, Tickets, KB Management"],
        ],
        widths=[1.4, 2.0, 3.1],
    )


def ai_pipeline() -> None:
    add_h1("7. AI Help and Ticket Pipeline")
    add_body(
        "The AI Help module is designed to reduce unnecessary ticket creation. Users first describe a technical problem. The system checks common Knowledge Base entries. Only when no matching Knowledge Base answer is available does the AI model generate troubleshooting steps."
    )
    add_picture_block(
        ASSET_DIR / "ai_help_pipeline_diagram.png",
        "Figure 2. AI Help pipeline showing Knowledge Base search, AI fallback, user decision, chat history, ticket creation, and escalation.",
    )
    DOC.add_page_break()
    add_table(
        ["Step", "Action", "Stored Data"],
        [
            ["1", "User enters issue.", "Issue text and signed-in user context."],
            ["2", "Knowledge Base keyword search runs.", "Matched entry source if available."],
            ["3", "If no match exists, AI suggestion is generated.", "AI response text and source value."],
            ["4", "User selects Issue Resolved.", "Chat history with result: Resolved without ticket."],
            ["5", "User selects Human Support.", "Ticket record plus chat history linked to ticket ID."],
            ["6", "Ticket priority and team are inferred.", "Priority, assigned team, status Open, timestamps."],
        ],
        widths=[0.55, 3.0, 2.95],
    )
    add_h2("Priority Rules")
    add_table(
        ["Detected Issue Type", "Priority"],
        [
            ["Server down, system down, access denied, email blocked", "High"],
            ["Critical, security breach, data loss, ransomware", "Critical"],
            ["VPN, network, Wi-Fi, email, printer, software", "Medium"],
            ["Password reset or general low-impact issue", "Low"],
        ],
        widths=[4.8, 1.7],
    )
    add_h2("Team Assignment Rules")
    add_table(
        ["Issue Keywords", "Assigned Team"],
        [
            ["VPN, Network, Wi-Fi", "Network Team"],
            ["Email, Outlook, Mail", "IT Support Team"],
            ["Printer, Laptop, Hardware, Display", "Hardware Team"],
            ["Password, Login, Access, MFA", "Access Management Team"],
            ["Software, Application, Installation", "Software Support Team"],
            ["Unknown issue", "General Support Team"],
        ],
        widths=[3.5, 3.0],
    )


def database_design() -> None:
    add_h1("8. Supabase Database Design")
    add_body(
        "The live Supabase database has been repaired and reset using the project SQL script. The schema now supports application-owned users, ticket records, Knowledge Base entries, chat history, and engineer resolution history."
    )
    DOC.add_page_break()
    add_table(
        ["Table", "Main Purpose", "Important Fields"],
        [
            ["users", "Stores user, support engineer, and admin accounts.", "id, full_name, email, role, employee_id, department, password_hash, created_at"],
            ["knowledge_base", "Stores common IT issue solutions.", "id, issue_name, keywords, solution, created_at, updated_at"],
            ["tickets", "Stores support tickets and lifecycle status.", "ticket_id, user_id, user details, issue, AI suggestion, priority, assigned team, engineer, status, resolution, timestamps"],
            ["chat_history", "Stores every AI/Knowledge Base support interaction.", "user details, issue, ai_response, source, result, related ticket_id, created_at"],
            ["ticket_resolution_history", "Stores engineer workflow updates.", "ticket_id, engineer, previous_status, new_status, remarks, resolution_notes, created_at"],
        ],
        widths=[1.65, 2.0, 2.85],
    )
    add_h2("Current Verified Demo Data")
    add_table(
        ["Data Area", "Verified Status"],
        [
            ["Users", "4 accounts: 2 users, 1 support engineer, 1 admin."],
            ["Knowledge Base", "16 common IT issue entries."],
            ["Tickets", "2 clean demo tickets after reset."],
            ["Chat History", "2 demo AI/Knowledge Base interactions."],
            ["Resolution History", "1 engineer update linked to TKT-0001."],
        ],
        widths=[2.0, 4.5],
    )
    add_callout(
        "Database repair completed",
        "The earlier problem with an old auth.users foreign key and missing password_hash was fixed through the SQL repair script. The database repair script was executed and the repaired schema was verified successfully.",
    )


def ui_ux() -> None:
    add_h1("9. UI and UX Development")
    add_body(
        "The project was redesigned from a single-page content dump into a structured multi-page SaaS-style application. Each page now has a focused purpose, clear hierarchy, proper spacing, role-specific navigation, and relevant visual assets."
    )
    add_table(
        ["Page / Area", "UI Work Completed"],
        [
            ["Home", "Clean landing page with navbar, hero, feature cards, simple workflow, FAQ, and footer."],
            ["Login and Signup", "Role-based forms, support engineer department list, session redirection, and logout flow."],
            ["AI Help", "Aligned issue input, result panel, four action buttons, Knowledge Base answer, AI suggestion, and 3D visual."],
            ["Raise Ticket", "Submit Support Request heading, placeholders, clean form, latest ticket preview, no-ticket layout, and View All Tickets action."],
            ["My Tickets", "Paginated user ticket list with ticket ID, issue, priority, status, and View Detail button."],
            ["Knowledge Base", "Search bar, quick review issues, list layout, aligned issue type/category labels, and no keyword clutter."],
            ["Engineer Dashboard", "Analytics-style dashboard, KPI cards, ticket table, search/filter/sort, pagination, and ticket management."],
            ["Admin Dashboard", "Operational control dashboard, analytics cards, reports, users list, engineer list, and Knowledge Base management."],
        ],
        widths=[1.55, 4.95],
    )
    add_bullets(
        [
            "Project name changed to Imdad AI.",
            "Custom logo and watermark-style brand mark added across the project.",
            "Theme refined using a consistent SaaS visual direction inspired by the NEFA reference, adapted to the Imdad AI support domain.",
            "Scroll transitions were added and later smoothed to reduce UI lag.",
            "Navbar active highlighting was added so users know which page they are on.",
        ]
    )
    add_h2("Implementation Evidence Screenshots")
    add_body(
        "The following role-based collages show representative screens from the implemented application. They connect the documented workflow to visible application output."
    )
    DOC.add_page_break()
    add_picture_block(
        ASSET_DIR / "user_pages_collage.png",
        "Figure 3. User pages: AI Help, Raise Ticket, and My Tickets.",
        width=6.15,
    )
    DOC.add_page_break()
    add_picture_block(
        ASSET_DIR / "engineer_pages_collage.png",
        "Figure 4. Support Engineer pages: dashboard, ticket management, and ticket detail review.",
        width=6.15,
    )
    DOC.add_page_break()
    add_picture_block(
        ASSET_DIR / "admin_pages_collage.png",
        "Figure 5. Admin pages: analytics dashboard, Knowledge Base management, and ticket monitoring.",
        width=6.15,
    )


def engineer_workflow() -> None:
    add_h1("10. Engineer Workflow and Resolution History")
    add_body(
        "The Support Engineer module now goes beyond viewing tickets. It supports ticket operations, update tracking, and reporting behavior that resembles a real service desk workflow."
    )
    add_table(
        ["Engineer Feature", "Implementation Status"],
        [
            ["Login system", "Available with role-based support engineer login."],
            ["Session management", "Current session persists in browser storage and controls navigation."],
            ["Protected dashboard", "Engineer pages require the correct role."],
            ["KPI cards", "Dashboard shows total, open, in progress, resolved, and workload indicators."],
            ["Ticket statistics", "Status, priority, and category overview included."],
            ["Open ticket management", "Engineer can view, search, filter, and open ticket details."],
            ["AI review", "Ticket detail includes AI suggestion and issue context."],
            ["Resolution workflow", "Engineer can assign, update status, add remarks, add resolution notes, resolve, and close tickets."],
            ["Resolution history", "Every engineer update is saved in the ticket_resolution_history table."],
        ],
        widths=[2.2, 4.3],
    )


def admin_analytics() -> None:
    add_h1("11. Admin Analytics and Reporting")
    add_body(
        "The Admin dashboard has been shaped as the main analytics area, so the separate Analytics navigation item was removed. Admin Home now provides system control, ticket visibility, and reporting in one place."
    )
    add_table(
        ["Report Area", "What It Shows"],
        [
            ["KPI cards", "Total users, support engineers, total tickets, open tickets, resolved tickets, high-priority tickets."],
            ["Tickets by status", "Open, In Progress, Resolved, and Closed ticket distribution."],
            ["Tickets by priority", "Critical, High, Medium, and Low ticket distribution."],
            ["Category-wise analysis", "Ticket breakdown by support area such as Access, Hardware, Network, Software, and Email."],
            ["Resolution trends", "Resolved and closed ticket movement based on current ticket status and history."],
            ["Engineer performance", "Assigned workload and completed work per engineer."],
            ["Ticket analytics table", "Searchable, filterable, sortable, and paginated ticket report for admin review."],
        ],
        widths=[2.1, 4.4],
    )
    add_callout(
        "Realtime data behavior",
        "The app subscribes to Supabase data changes and also uses a timed refresh fallback, so dashboards remain current even if realtime publication behavior is delayed in the demo environment.",
    )


def testing() -> None:
    add_h1("12. Testing and Verification")
    add_body(
        "Testing focused on the full support lifecycle: login, role redirection, Knowledge Base search, AI fallback, ticket creation, ticket visibility, engineer workflow, resolution history, admin analytics, and Supabase data persistence."
    )
    add_table(
        ["Test Case", "Expected Result", "Actual Result", "Status"],
        [
            ["Build verification", "TypeScript and Vite build should complete without errors.", "`npm run build` completed successfully after the latest implementation changes.", "PASS"],
            ["User login", "User should access only user-specific pages after login.", "User role opened AI Help, Raise Ticket, My Tickets, and Knowledge Base views.", "PASS"],
            ["Support engineer login", "Engineer should access protected engineer dashboard and ticket tools.", "Support Engineer role opened dashboard, ticket queue, and ticket detail workflow.", "PASS"],
            ["Admin login", "Admin should access system monitoring and management pages.", "Admin role opened analytics, users, tickets, and Knowledge Base management views.", "PASS"],
            ["Ticket creation", "Unresolved issue should create a structured ticket with generated metadata.", "Ticket was created with ticket ID, priority, team, status, timestamps, and AI context.", "PASS"],
            ["My Tickets access", "User should see only tickets linked to the signed-in account.", "User ticket list displayed account-linked tickets with pagination and detail action.", "PASS"],
            ["Engineer update", "Engineer should update assignment, status, remarks, and resolution notes.", "Engineer workflow saved assignment/status updates and resolution notes.", "PASS"],
            ["Resolution history", "Engineer updates should be recorded in resolution history.", "Updates were saved in `ticket_resolution_history` and shown on ticket details.", "PASS"],
            ["Supabase reset", "Demo reset should recreate clean users, KB entries, tickets, chat history, and history rows.", "Repair/reset script produced the expected clean demo dataset.", "PASS"],
            ["Live data check", "Repaired users table should support insert/delete with password_hash.", "Supabase insert/delete verification succeeded on the repaired schema.", "PASS"],
        ],
        widths=[1.45, 2.05, 2.35, 0.65],
    )
    add_h2("Demo Commands")
    add_table(
        ["Command", "Purpose"],
        [
            ["npm run dev", "Starts the Vite frontend at 127.0.0.1:5173."],
            ["npm run server", "Starts the Express AI API at 127.0.0.1:8787."],
            ["npm run build", "Checks TypeScript and produces a production build."],
        ],
        widths=[2.1, 4.4],
    )


def challenges() -> None:
    add_h1("13. Challenges Faced and Solutions")
    add_table(
        ["Challenge", "Solution Applied"],
        [
            ["Supabase schema mismatch", "Updated SQL to match actual table columns such as issue_name and integer KB IDs."],
            ["Old auth.users foreign key", "Repaired users table to use project-owned UUID accounts instead of depending on deleted auth records."],
            ["Missing password_hash", "Added password_hash column and updated account persistence logic."],
            ["Ticket ID compatibility", "Updated generator to support existing legacy TCK IDs and new TKT IDs."],
            ["AI response was too generic for custom issues", "Improved fallback troubleshooting logic for slow systems, network, VPN, email, and general issues."],
            ["Single-page content overload", "Restructured into proper routes and role-specific pages."],
            ["UI alignment issues", "Iteratively refined navbar, AI Help, Raise Ticket, My Tickets, Knowledge Base, Engineer, Admin, and footer layouts."],
            ["Realtime reliability", "Added refresh fallback alongside Supabase realtime subscription."],
        ],
        widths=[2.35, 4.15],
    )


def milestone() -> None:
    add_h1("14. Current Completion Milestone")
    add_body(
        "The project has reached the stage where the main end-to-end prototype can be demonstrated to a mentor: users can request help, tickets are created and stored, engineers can manage tickets, admins can monitor analytics, and Supabase contains clean live demo data."
    )
    add_table(
        ["Module", "Status", "Notes"],
        [
            ["User Authentication and Dashboard", "Completed", "Signup, login, session, logout, role redirection, and user pages are implemented."],
            ["AI Help Module", "Completed", "Knowledge Base search, AI suggestion, resolved flow, and ticket escalation are implemented."],
            ["Ticket Creation and User Tracking", "Completed", "Automatic ticket ID, priority, team, status, My Tickets, details, and chat history are implemented."],
            ["Knowledge Base", "Completed", "User-facing search and admin management are implemented."],
            ["Engineer Authentication", "Completed", "Engineer login, protected dashboard, session, and logout are available."],
            ["Engineer Dashboard", "Completed", "KPIs, ticket stats, recent activity, priority/category overview are available."],
            ["Open Ticket Management", "Completed", "Search, filters, detail page, AI review, and recommendation review are available."],
            ["Resolution Workflow", "Completed", "Remarks, notes, status update, assignment, close ticket, and resolution history are available."],
            ["Reports and Analytics", "Completed", "Admin/engineer analytics, category, priority, trend, and performance metrics are available."],
            ["Production hardening", "Pending", "Needs Supabase Auth, stricter RLS, deployment, notifications, and audit logging."],
        ],
        widths=[2.2, 1.2, 3.1],
    )


def future_scope() -> None:
    add_h1("15. Future Scope and Presentation Notes")
    add_h2("Recommended Future Enhancements")
    add_bullets(
        [
            "Replace prototype role login with full Supabase Auth and secure role claims.",
            "Add strict Row Level Security policies so users can only read their own rows at the database level.",
            "Deploy frontend and backend to a cloud platform.",
            "Add email or Teams notifications when tickets are created, assigned, resolved, or closed.",
            "Add SLA timers, escalation rules, and overdue ticket alerts.",
            "Add file attachments for screenshots and error logs.",
            "Add downloadable reports for admin analytics.",
            "Add audit logging for all sensitive support operations.",
        ]
    )
    add_h2("How to Present the Project")
    add_table(
        ["Step", "Presentation Action"],
        [
            ["1", "Start with the problem: repeated IT issues waste support time and delay users."],
            ["2", "Introduce Imdad AI as the solution: first self-help, then AI, then ticket escalation."],
            ["3", "Show the Home page and role-based navbar."],
            ["4", "Log in as a User and demonstrate AI Help, Knowledge Base result, unresolved issue, and ticket creation."],
            ["5", "Open My Tickets and Ticket Details to show user-side tracking."],
            ["6", "Log in as Support Engineer and show dashboard, open tickets, AI suggestion review, status update, remarks, and resolution history."],
            ["7", "Log in as Admin and show analytics, users, engineers, tickets, and Knowledge Base management."],
            ["8", "End with the database: show Supabase tables proving the system stores real records."],
        ],
        widths=[0.7, 5.8],
    )
    add_callout(
        "Conclusion",
        "Imdad AI is a complete role-based AI support desk prototype with real database persistence, structured escalation, engineer workflow, and admin reporting. It demonstrates both product thinking and technical implementation from UI to database.",
    )


def appendix() -> None:
    add_h1("Appendix A: File and Database Reference")
    add_table(
        ["Project Area", "Important Files"],
        [
            ["Frontend app", "src/App.tsx, src/styles.css, src/main.tsx"],
            ["Supabase client and store", "src/supabaseClient.ts, src/supabaseStore.ts"],
            ["AI backend", "server/index.js"],
            ["Database scripts", "supabase/schema.sql, supabase/fix_schema_and_reset_demo.sql"],
            ["Brand and visuals", "public/brand, public/visuals"],
            ["Documentation", "docs folder"],
        ],
        widths=[2.0, 4.5],
    )
    add_h2("Clean Demo Records")
    add_table(
        ["Record Type", "Examples"],
        [
            ["Users", "Aarav Mehta, Neha Kapoor, Priya Nair, Maya Rao"],
            ["Tickets", "TKT-0001 Access denied to finance dashboard; TKT-0002 Laptop slow after update"],
            ["Knowledge Base topics", "VPN, Wi-Fi, email, password reset, access denied, slow system, printer, software, server down, file missing"],
            ["Resolution history", "TKT-0001 update from Open to In Progress by Priya Nair"],
        ],
        widths=[1.8, 4.7],
    )


cover_page()
contents_page()
executive_summary()
problem_objectives()
journey()
tech_stack()
architecture()
role_modules()
ai_pipeline()
database_design()
ui_ux()
engineer_workflow()
admin_analytics()
testing()
challenges()
milestone()
future_scope()
appendix()

for section in DOC.sections:
    set_footer(section)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
DOC.save(OUTPUT)
print(OUTPUT)
